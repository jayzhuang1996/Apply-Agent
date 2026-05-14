from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Formatting
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

doc.add_paragraph("Dear Hiring Manager,")

doc.add_paragraph(
    "I am writing to express my strong interest in the open position. You might notice something unique about this submission: this application was completed and submitted entirely by an autonomous AI agent that I architected and built from scratch. I designed this system to showcase my ability to solve complex operational bottlenecks through full-stack engineering and applied artificial intelligence."
)

doc.add_heading("The Core Engine", level=2)
doc.add_paragraph(
    "The traditional job application process is highly repetitive, requiring candidates to manually enter the same data into complex Applicant Tracking Systems (ATS). My vision was to fulfill the original prompt: build a completely end-to-end, fully autonomous agent that actively parses the job description, reasons about unstructured form fields, dynamically tailors my resume, and hits submit without any human intervention."
)
doc.add_paragraph(
    "This engine is powered by a modern tech stack. The backend runs on Python and FastAPI, utilizing background threading and queues for persistent automation. Playwright handles deterministic DOM manipulation, bypassing complex React components and shadow DOMs using JavaScript injection. Most importantly, I leverage OpenAI's language models as a semantic reasoning engine to map raw HTML form labels to structured responses."
)

doc.add_heading("The Creative Spin-off: Dual-Mode Execution", level=2)
doc.add_paragraph(
    "While the agent successfully navigates multi-page applications and autonomously submits highly targeted applications blind, I wanted to take the architecture to another level. I packaged the system with two distinct execution modes, transforming a simple script into a robust, open-source framework."
)
doc.add_paragraph(
    "First, I built an interactive 'Human-in-the-Loop' web dashboard. This layer streams live execution logs via Server-Sent Events, pauses on the final screen, and allows a user to instantly edit form fields via a bidirectional connection before submitting. This proves the system's viability in high-stakes enterprise scenarios where AI does the heavy lifting but requires human oversight."
)
doc.add_paragraph(
    "Second, I packaged the entire architecture into an 'AI Skill' for agentic coding IDEs like Claude Code. By open-sourcing the repository with a dedicated SKILL.md file, any developer can install the tool and instruct their own local AI to autonomously drive the background Python scripts. It turns the codebase into an executable tool that an AI can use natively in a chat window."
)

doc.add_heading("Hardening the Agent: Bypassing Bot Detection", level=2)
doc.add_paragraph(
    "One of the biggest technical hurdles in modern web automation is the proliferation of anti-bot solutions like Cloudflare Turnstile. To ensure the agent's reliability, I integrated 'playwright-stealth' to mask automation signatures, such as the webdriver property and specialized navigator objects. Furthermore, I designed the system to be 'hybrid-aware'—if a CAPTCHA is encountered, the agent can signal for manual intervention in the visible browser window, seamlessly combining AI efficiency with human problem-solving."
)

doc.add_heading("Conclusion", level=2)
doc.add_paragraph(
    "Building this end-to-end system—from the autonomous web scraper to the reactive UI and the agentic skill package—demonstrates my deep technical execution, systems thinking, and passion for automation. I look forward to bringing this same builder's mindset and drive for efficiency to your team."
)

doc.add_paragraph("\nSincerely,\nJay Zhuang")

doc.save("output/Cover_Letter_Agent.docx")
print("Cover letter updated at output/Cover_Letter_Agent.docx")
